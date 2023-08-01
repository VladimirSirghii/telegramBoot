from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Crearea documentului
document = Document()

# Adăugarea titlului și numărului de laborator
document.add_heading('Microsoft Word', level=1)
document.add_heading('Lucrare de laborator nr. 4', level=2)

# Salvarea documentului pe unitatea D: în dosarul personal
file_path = r'C:\Users\vovci\OneDrive\Desktop\Lucrari individuala\Varzari_Nina.docx'
document.save(file_path)

# Adăugarea antetului cu specialitatea
header = document.sections[0].header
header_paragraph = header.paragraphs[0]
header_paragraph.text = 'Specialitatea dvs.'
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Adăugarea secțiunii SIMBOLURI
document.add_heading('SIMBOLURI', level=2)

# Adăugarea simbolurilor utilizând fontul Wingdings
symbols = ['#', '§', '®', '∞', '↔', '×', 'Ɏ', '˄', '˅', 'Ω', 'Ψ', 'Δ', 'Σ', 'Φ', 'α', 'β', 'γ', 'δ', 'ε', '⁂',
           '☺', '☻', '☼', '♫', '', '', '', '', '', '', '', '', '']
symbol_paragraph = document.add_paragraph()
symbol_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for symbol in symbols:
    run = symbol_paragraph.add_run(symbol)
    run.font.name = 'Wingdings'

# Adăugarea secțiunii FORMULE
document.add_heading('FORMULE', level=2)

# Adăugarea formulelor utilizând comanda Equation
formulas = ['Economie circulară', 'Producție', 'Uz / Consum', 'Reciclare Deșeuri', 'Resurse',
            'Interdependența concurență - eficiență - progres economic']
formula_paragraph = document.add_paragraph()
formula_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for formula in formulas:
    formula_paragraph.add_run(formula)
    formula_paragraph.add_run('\n')

# Adăugarea secțiunii DIAGRAME
document.add_heading('DIAGRAME', level=2)

# Adăugarea titlurilor și realizarea diagramelor
diagram_titles = ['Diagrama 1', 'Diagrama 2']
diagram_descriptions = ['Economie circulară\n\nProducție\n\nUz / Consum\n\nReciclare Deșeuri\n\nResurse',
                        'Interdependența concurență - eficiență - progres economic']
for i in range(len(diagram_titles)):
    document.add_heading(diagram_titles[i], level=3)
    document.add_paragraph(diagram_descriptions[i])

# Adăugarea secțiunii IMAGINI
document.add_heading('IMAGINI', level=2)

# Adăugarea imaginilor (necesită imagini salvate pe calculator)
image_paths = [
    r'C:\Users\vovci\OneDrive\Desktop\Lucrari individuala\Screenshot_1.png',
    r'C:\Users\vovci\OneDrive\Desktop\Lucrari individuala\Screenshot_1.png',
    r'C:\Users\vovci\OneDrive\Desktop\Lucrari individuala\Screenshot_1.png',
    r'C:\Users\vovci\OneDrive\Desktop\Lucrari individuala\Screenshot_1.png',
]
for image_path in image_paths:
    document.add_picture(image_path)

# Adăugarea numerotației paginilor în partea de jos, centrat
footer = document.sections[0].footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = 'Pagina {PAGE}'
footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Aplicarea stilurilor Heading pentru titluri
heading_styles = ['Heading1', 'Heading2', 'Heading2', 'Heading2']
headings = ['SIMBOLURI', 'FORMULE', 'DIAGRAME', 'IMAGINI']

for i in range(len(heading_styles)):
    heading = document.add_heading(level=i+1)
    heading.style = heading_styles[i]
    heading.text = headings[i]

# Inserarea întreruperii de pagină pentru a trece la o nouă pagină
document.add_page_break()

# Adăugarea cuvântului CUPRINS pe o pagină nouă
document.add_paragraph().text = 'CUPRINS'
paragraph = document.paragraphs[-1]
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
paragraph.runs[0].bold = True
paragraph.runs[0].font.size = Pt(14)
document.add_paragraph().text = '\n'

# Inserarea cuprinsului automat
document.add_paragraph().text = 'Table of Contents'

# Salvarea finală a documentului
document.save(file_path)
